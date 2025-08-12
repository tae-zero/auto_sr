from fastapi import APIRouter, HTTPException, Request, UploadFile, File, Form
from typing import Dict, List, Optional, Any
import logging
import os
import json
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
import tempfile
import shutil

logger = logging.getLogger(__name__)

router = APIRouter()

# AI 서비스 설정
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")

# 벡터 데이터베이스 초기화
vector_store = None
conversation_chain = None

def initialize_ai_services():
    """AI 서비스 초기화"""
    global vector_store, conversation_chain
    
    if not OPENAI_API_KEY:
        logger.warning("⚠️ OPENAI_API_KEY가 설정되지 않았습니다.")
        return False
    
    try:
        # OpenAI 모델 초기화
        llm = ChatOpenAI(
            openai_api_key=OPENAI_API_KEY,
            model_name=OPENAI_MODEL,
            temperature=0.7
        )
        
        # 임베딩 모델 초기화
        embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
        
        # 벡터 데이터베이스 초기화
        persist_directory = os.getenv("CHROMA_PERSIST_DIRECTORY", "./chroma_db")
        if os.path.exists(persist_directory):
            vector_store = Chroma(
                persist_directory=persist_directory,
                embedding_function=embeddings
            )
            logger.info("✅ 기존 벡터 데이터베이스 로드 완료")
        else:
            # 빈 벡터 스토어 생성
            vector_store = Chroma(
                embedding_function=embeddings,
                persist_directory=persist_directory
            )
            logger.info("✅ 새로운 벡터 데이터베이스 생성 완료")
        
        # 대화 체인 초기화
        memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True
        )
        
        conversation_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=vector_store.as_retriever(),
            memory=memory,
            return_source_documents=True
        )
        
        logger.info("✅ AI 서비스 초기화 완료")
        return True
        
    except Exception as e:
        logger.error(f"❌ AI 서비스 초기화 실패: {str(e)}")
        return False

@router.on_event("startup")
async def startup_event():
    """서비스 시작 시 AI 서비스 초기화"""
    initialize_ai_services()

@router.get("/ai/health",
    summary="AI 서비스 상태 확인",
    description="AI 서비스와 LangChain의 상태를 확인합니다.",
    response_description="AI 서비스 상태",
    tags=["AI 서비스"]
)
async def ai_health_check():
    """AI 서비스 health check"""
    try:
        if not OPENAI_API_KEY:
            return {
                "status": "unhealthy",
                "message": "OPENAI_API_KEY가 설정되지 않았습니다.",
                "ai_services": "disabled"
            }
        
        if vector_store and conversation_chain:
            return {
                "status": "healthy",
                "message": "AI 서비스가 정상적으로 작동 중입니다.",
                "ai_services": "enabled",
                "vector_db": "connected",
                "llm_model": OPENAI_MODEL
            }
        else:
            return {
                "status": "unhealthy",
                "message": "AI 서비스 초기화에 실패했습니다.",
                "ai_services": "failed"
            }
    except Exception as e:
        logger.error(f"AI 서비스 health check 실패: {str(e)}")
        return {
            "status": "error",
            "message": f"AI 서비스 오류: {str(e)}",
            "ai_services": "error"
        }

@router.post("/ai/chat",
    summary="AI 챗봇 대화",
    description="LangChain을 사용하여 AI와 대화합니다.",
    response_description="AI 응답",
    tags=["AI 서비스"]
)
async def ai_chat(
    message: str = Form(...),
    chat_history: Optional[str] = Form("[]")
):
    """AI 챗봇 대화"""
    try:
        if not conversation_chain:
            raise HTTPException(status_code=503, detail="AI 서비스가 초기화되지 않았습니다.")
        
        # 채팅 히스토리 파싱
        try:
            history = json.loads(chat_history) if chat_history else []
        except:
            history = []
        
        # AI 응답 생성
        response = conversation_chain.invoke({
            "question": message,
            "chat_history": history
        })
        
        return {
            "response": response["answer"],
            "sources": [doc.page_content[:100] + "..." for doc in response.get("source_documents", [])],
            "chat_history": history + [{"user": message, "ai": response["answer"]}]
        }
        
    except Exception as e:
        logger.error(f"AI 챗봇 대화 실패: {str(e)}")
        raise HTTPException(status_code=500, detail=f"AI 챗봇 오류: {str(e)}")

@router.post("/ai/upload-document",
    summary="문서 업로드 및 벡터화",
    description="문서를 업로드하여 벡터 데이터베이스에 저장합니다.",
    response_description="업로드 결과",
    tags=["AI 서비스"]
)
async def upload_document(
    file: UploadFile = File(...),
    description: Optional[str] = Form("")
):
    """문서 업로드 및 벡터화"""
    try:
        if not vector_store:
            raise HTTPException(status_code=503, detail="AI 서비스가 초기화되지 않았습니다.")
        
        # 파일 크기 검증
        max_size = int(os.getenv("MAX_FILE_SIZE", "10485760"))  # 10MB
        if file.size > max_size:
            raise HTTPException(status_code=400, detail=f"파일 크기가 너무 큽니다. 최대 {max_size/1024/1024}MB")
        
        # 파일 타입 검증
        supported_types = os.getenv("SUPPORTED_FILE_TYPES", "txt,pdf,docx,md").split(",")
        file_extension = file.filename.split(".")[-1].lower()
        if file_extension not in supported_types:
            raise HTTPException(status_code=400, detail=f"지원하지 않는 파일 타입입니다. 지원: {', '.join(supported_types)}")
        
        # 임시 파일로 저장
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
            shutil.copyfileobj(file.file, temp_file)
            temp_path = temp_file.name
        
        try:
            # 문서 내용 추출 및 벡터화
            documents = await process_document(temp_path, file_extension, description)
            
            # 벡터 데이터베이스에 저장
            vector_store.add_documents(documents)
            vector_store.persist()
            
            return {
                "message": "문서가 성공적으로 업로드되었습니다.",
                "filename": file.filename,
                "documents_count": len(documents),
                "description": description
            }
            
        finally:
            # 임시 파일 삭제
            os.unlink(temp_path)
            
    except Exception as e:
        logger.error(f"문서 업로드 실패: {str(e)}")
        raise HTTPException(status_code=500, detail=f"문서 업로드 오류: {str(e)}")

async def process_document(file_path: str, file_type: str, description: str = "") -> List[Document]:
    """문서 처리 및 텍스트 추출"""
    try:
        if file_type == "txt":
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif file_type == "md":
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif file_type == "pdf":
            import PyPDF2
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
        elif file_type == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        else:
            raise ValueError(f"지원하지 않는 파일 타입: {file_type}")
        
        # 텍스트 분할
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        
        chunks = text_splitter.split_text(text)
        
        # Document 객체 생성
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "chunk_id": i,
                    "description": description,
                    "file_type": file_type
                }
            )
            documents.append(doc)
        
        return documents
        
    except Exception as e:
        logger.error(f"문서 처리 실패: {str(e)}")
        raise

@router.get("/ai/search",
    summary="벡터 검색",
    description="벡터 데이터베이스에서 유사한 문서를 검색합니다.",
    response_description="검색 결과",
    tags=["AI 서비스"]
)
async def vector_search(query: str, top_k: int = 5):
    """벡터 검색"""
    try:
        if not vector_store:
            raise HTTPException(status_code=503, detail="AI 서비스가 초기화되지 않았습니다.")
        
        # 벡터 검색 수행
        results = vector_store.similarity_search(query, k=top_k)
        
        return {
            "query": query,
            "results": [
                {
                    "content": doc.page_content,
                    "metadata": doc.metadata
                }
                for doc in results
            ],
            "total_results": len(results)
        }
        
    except Exception as e:
        logger.error(f"벡터 검색 실패: {str(e)}")
        raise HTTPException(status_code=500, detail=f"검색 오류: {str(e)}")
