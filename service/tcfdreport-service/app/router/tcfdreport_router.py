from fastapi import APIRouter, HTTPException, Depends, BackgroundTasks
from fastapi.responses import FileResponse, StreamingResponse, HTMLResponse
from typing import Dict, Any, List
import logging
import asyncpg
import os
from datetime import datetime
import tempfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import zipfile
import urllib.parse
from weasyprint import HTML, CSS

from app.domain.tcfd.entity.tcfd_input_entity import TCFDInputEntity
from app.domain.tcfd.entity.tcfd_draft_entity import TCFDDraftEntity
from app.domain.tcfd.schema.tcfd_input_schema import TCFDInputCreateSchema, TCFDInputUpdateSchema, TCFDInputResponseSchema
from app.domain.tcfd.schema.tcfd_draft_schema import TCFDDraftCreateSchema, TCFDDraftUpdateSchema, TCFDDraftResponseSchema
from app.domain.tcfd.repository.tcfd_input_repository import TCFDInputRepository
from app.domain.tcfd.repository.tcfd_draft_repository import TCFDDraftRepository

logger = logging.getLogger(__name__)

tcfdreport_router = APIRouter()

def cleanup_temp_files(*file_paths: str):
    """임시 파일들을 정리하는 함수"""
    for file_path in file_paths:
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
                logger.info(f"✅ 임시 파일 정리 완료: {file_path}")
        except Exception as e:
            logger.warning(f"⚠️ 임시 파일 정리 실패: {file_path} - {e}")

async def _return_html_fallback(data: Dict[str, Any], error_type: str = "unknown"):
    """HTML fallback 반환 (WeasyPrint 오류 시)"""
    try:
        company_name = data.get('company_name', 'TCFD')
        if company_name == 'TCFD' and data.get('draft'):
            draft_content = data['draft']
            if '**회사명**:' in draft_content:
                company_name = draft_content.split('**회사명**:')[1].split('\n')[0].strip()
            elif '회사명:' in draft_content:
                company_name = draft_content.split('회사명:')[1].split('\n')[0].strip()
        
        safe_company_name = company_name.replace('*', '').replace('/', '_').replace('\\', '_').replace(':', '_').replace('|', '_').replace('<', '_').replace('>', '_').replace('"', '_').replace('?', '_')
        if len(safe_company_name) > 20:
            safe_company_name = safe_company_name[:20]
        
        filename = f"{safe_company_name}_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{error_type}.html"
        filename_encoded = urllib.parse.quote(filename)
        
        # HTML 콘텐츠 생성
        html_content = f"""
        <!DOCTYPE html>
        <html lang="ko">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{company_name} TCFD 보고서</title>
            <style>
                body {{ 
                    font-family: 'Noto Sans KR', 'Malgun Gothic', Arial, sans-serif; 
                    line-height: 1.6;
                    margin: 0;
                    padding: 20px;
                    background-color: #f8f9fa;
                }}
                .container {{
                    max-width: 800px;
                    margin: 0 auto;
                    background: white;
                    padding: 30px;
                    border-radius: 10px;
                    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
                }}
                h1 {{ 
                    text-align: center; 
                    color: #2563eb; 
                    border-bottom: 2px solid #2563eb; 
                    padding-bottom: 10px; 
                }}
                h2 {{ 
                    color: #059669; 
                    margin-top: 30px; 
                }}
                .company-info {{ 
                    text-align: center; 
                    color: #6b7280; 
                    margin: 20px 0; 
                    padding: 15px;
                    background: #f1f5f9;
                    border-radius: 8px;
                }}
                .content {{ 
                    background: #f9fafb; 
                    padding: 20px; 
                    border-radius: 8px; 
                    margin: 20px 0; 
                    border-left: 4px solid #3b82f6;
                }}
                .error-notice {{
                    background: #fef3c7;
                    border: 1px solid #f59e0b;
                    color: #92400e;
                    padding: 15px;
                    border-radius: 8px;
                    margin: 20px 0;
                    text-align: center;
                }}
                .timestamp {{ 
                    text-align: center; 
                    color: #9ca3af; 
                    font-size: 14px; 
                    margin: 20px 0; 
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <h1>{company_name} TCFD 보고서</h1>
                <div class="company-info">생성일시: {datetime.now().strftime("%Y년 %m월 %d일 %H시 %M분")}</div>
                
                <div class="error-notice">
                    PDF 생성 중 오류가 발생하여 HTML 형태로 제공됩니다.<br>
                    오류 유형: {error_type}
                </div>
                
                <h2>AI 생성 초안</h2>
                <div class="content">{data.get('draft', '').replace(chr(10), '<br>')}</div>
                
                <h2>윤문된 텍스트</h2>
                <div class="content">{data.get('polished', '').replace(chr(10), '<br>')}</div>
                
                <div class="timestamp">
                    이 문서는 {datetime.now().strftime("%Y년 %m월 %d일 %H시 %M분")}에 생성되었습니다.
                </div>
            </div>
        </body>
        </html>
        """
        
        logger.info(f"HTML fallback 반환: {filename}")
        
        # StreamingResponse로 반환
        response = StreamingResponse(
            io.BytesIO(html_content.encode('utf-8')),
            media_type="text/html",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{filename_encoded}",
                "Cache-Control": "no-cache, no-store, must-revalidate",
                "Pragma": "no-cache",
                "Expires": "0",
                "X-Content-Type-Options": "nosniff",
                "X-Frame-Options": "DENY"
            }
        )
        
        return response
        
    except Exception as e:
        logger.error(f"HTML fallback 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=f"HTML fallback 생성 실패: {str(e)}")

# 데이터베이스 연결 함수
async def get_db_connection():
    """데이터베이스 연결 반환"""
    try:
        database_url = os.getenv("DATABASE_URL")
        if not database_url:
            raise Exception("DATABASE_URL 환경변수가 설정되지 않았습니다")
        
        # URL 스키마 수정
        if database_url.startswith("postgresql+asyncpg://"):
            database_url = "postgresql://" + database_url[len("postgresql+asyncpg://"):]
        elif database_url.startswith("postgres://"):
            database_url = "postgresql://" + database_url[len("postgres://"):]
        
        conn = await asyncpg.connect(database_url)
        return conn
    except Exception as e:
        logger.error(f"데이터베이스 연결 실패: {str(e)}")
        raise

@tcfdreport_router.get("/")
async def root():
    return {"message": "TCFD Report Service"}

@tcfdreport_router.get("/health")
async def health_check():
    return {"status": "healthy"}

@tcfdreport_router.post("/inputs")
async def create_tcfd_inputs(data: Dict[str, Any]):
    """TCFD 입력 데이터 생성"""
    try:
        conn = await get_db_connection()
        
        # 실제 테이블 구조에 맞춰 INSERT 쿼리 수정
        result = await conn.fetchrow(
            """
            INSERT INTO tcfd_inputs (
                company_name, user_id,
                governance_g1, governance_g2,
                strategy_s1, strategy_s2, strategy_s3,
                risk_management_r1, risk_management_r2, risk_management_r3,
                metrics_targets_m1, metrics_targets_m2, metrics_targets_m3,
                created_at, updated_at
            )
            VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15)
            RETURNING *
            """,
            data.get('company_name'),
            data.get('user_id'),
            data.get('governance_g1'),
            data.get('governance_g2'),
            data.get('strategy_s1'),
            data.get('strategy_s2'),
            data.get('strategy_s3'),
            data.get('risk_management_r1'),
            data.get('risk_management_r2'),
            data.get('risk_management_r3'),
            data.get('metrics_targets_m1'),
            data.get('metrics_targets_m2'),
            data.get('metrics_targets_m3'),
            datetime.now(),
            datetime.now()
        )
        
        if result:
            # Record를 dict로 변환
            result_dict = dict(result)
            await conn.close()
            return {"success": True, "data": result_dict}
        else:
            await conn.close()
            raise HTTPException(status_code=500, detail="데이터 삽입 실패")
            
    except Exception as e:
        logger.error(f"TCFD 입력 데이터 생성 실패: {str(e)}")
        raise HTTPException(status_code=500, detail=f"TCFD 입력 데이터 생성 실패: {str(e)}")

@tcfdreport_router.get("/inputs/{company_name}")
async def get_tcfd_inputs(company_name: str):
    """회사별 TCFD 데이터 조회"""
    try:
        conn = await get_db_connection()
        
        # asyncpg Record를 dict로 캐스팅
        result = await conn.fetchrow(
            """
            SELECT * FROM tcfd_inputs 
            WHERE company_name = $1 
            ORDER BY created_at DESC 
            LIMIT 1
            """,
            company_name
        )
        
        if result:
            # Record를 dict로 변환
            result_dict = dict(result)
            await conn.close()
            return {"success": True, "data": result_dict}
        else:
            await conn.close()
            return {"success": False, "message": "데이터를 찾을 수 없습니다"}
            
    except Exception as e:
        logger.error(f"TCFD 데이터 조회 실패: {str(e)}")
        raise HTTPException(status_code=500, detail=f"TCFD 데이터 조회 실패: {str(e)}")

@tcfdreport_router.post("/drafts")
async def create_tcfd_draft(data: TCFDDraftCreateSchema):
    """TCFD 초안 데이터 생성"""
    try:
        conn = await get_db_connection()
        
        # Entity 생성
        draft_entity = TCFDDraftEntity(
            company_name=data.company_name,
            user_id=data.user_id,
            tcfd_input_id=data.tcfd_input_id,
            draft_content=data.draft_content,
            draft_type=data.draft_type,
            file_path=data.file_path,
            status=data.status
        )
        
        # Repository를 통해 저장
        repository = TCFDDraftRepository()
        saved_draft = await repository.save(conn, draft_entity)
        
        await conn.close()
        return {"success": True, "data": saved_draft.to_dict()}
        
    except Exception as e:
        logger.error(f"TCFD 초안 데이터 생성 실패: {str(e)}")
        raise HTTPException(status_code=500, detail=f"TCFD 초안 데이터 생성 실패: {str(e)}")

@tcfdreport_router.get("/drafts/{company_name}")
async def get_tcfd_drafts(company_name: str):
    """회사별 TCFD 초안 데이터 조회"""
    try:
        conn = await get_db_connection()
        
        repository = TCFDDraftRepository()
        drafts = await repository.find_by_company_name(conn, company_name)
        
        await conn.close()
        
        drafts_data = [draft.to_dict() for draft in drafts]
        return {"success": True, "data": drafts_data, "total_count": len(drafts_data)}
        
    except Exception as e:
        logger.error(f"TCFD 초안 데이터 조회 실패: {str(e)}")
        raise HTTPException(status_code=500, detail=f"TCFD 초안 데이터 조회 실패: {str(e)}")

@tcfdreport_router.get("/drafts/id/{draft_id}")
async def get_tcfd_draft_by_id(draft_id: int):
    """ID로 TCFD 초안 데이터 조회"""
    try:
        conn = await get_db_connection()
        
        repository = TCFDDraftRepository()
        draft = await repository.find_by_id(conn, draft_id)
        
        await conn.close()
        
        if draft:
            return {"success": True, "data": draft.to_dict()}
        else:
            raise HTTPException(status_code=404, detail="초안 데이터를 찾을 수 없습니다")
        
    except Exception as e:
        logger.error(f"TCFD 초안 데이터 조회 실패: {str(e)}")
        raise HTTPException(status_code=500, detail=f"TCFD 초안 데이터 조회 실패: {str(e)}")

@tcfdreport_router.put("/drafts/{draft_id}/status")
async def update_draft_status(draft_id: int, status: str):
    """TCFD 초안 데이터 상태 업데이트"""
    try:
        conn = await get_db_connection()
        
        repository = TCFDDraftRepository()
        success = await repository.update_status(conn, draft_id, status)
        
        await conn.close()
        
        if success:
            return {"success": True, "message": "상태 업데이트 완료"}
        else:
            raise HTTPException(status_code=404, detail="초안 데이터를 찾을 수 없습니다")
        
    except Exception as e:
        logger.error(f"TCFD 초안 데이터 상태 업데이트 실패: {str(e)}")
        raise HTTPException(status_code=500, detail=f"TCFD 초안 데이터 상태 업데이트 실패: {str(e)}")

@tcfdreport_router.post("/download/word")
async def download_tcfd_report_as_word(data: Dict[str, Any]):
    """TCFD 보고서를 Word로 다운로드"""
    try:
        logger.info(f"Word 다운로드 요청: {data.get('company_name', 'Unknown')}")
        
        # 회사명 추출
        company_name = data.get('company_name', 'TCFD')
        if company_name == 'TCFD' and data.get('draft'):
            draft_content = data['draft']
            if '**회사명**:' in draft_content:
                company_name = draft_content.split('**회사명**:')[1].split('\n')[0].strip()
            elif '회사명:' in draft_content:
                company_name = draft_content.split('회사명:')[1].split('\n')[0].strip()
        
        # 파일명에 사용할 수 없는 특수문자 제거 (한글은 유지)
        safe_company_name = company_name.replace('*', '').replace('/', '_').replace('\\', '_').replace(':', '_').replace('|', '_').replace('<', '_').replace('>', '_').replace('"', '_').replace('?', '_')
        if len(safe_company_name) > 20:
            safe_company_name = safe_company_name[:20]
        
        # Word 문서 생성
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading(f'{company_name} TCFD 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 회사 정보 추가
        doc.add_paragraph(f'생성일시: {datetime.now().strftime("%Y년 %m월 %d일 %H시 %M분")}')
        doc.add_paragraph('')
        
        # 초안 내용 추가
        doc.add_heading('AI 생성 초안', level=1)
        doc.add_paragraph(data.get('draft', ''))
        
        # 윤문된 텍스트 추가
        doc.add_heading('윤문된 텍스트', level=1)
        doc.add_paragraph(data.get('polished', ''))
        
        # 파일명 생성 (한글 포함)
        filename = f"{safe_company_name}_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # 메모리에 저장
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info(f"Word 문서 생성 성공: {filename}, size={len(doc_bytes.getvalue())}B")
        
        # 한글 파일명을 위한 안전한 헤더 설정
        try:
            # UTF-8로 인코딩 시도
            filename_encoded = urllib.parse.quote(filename, safe='')
            content_disposition = f"attachment; filename*=UTF-8''{filename_encoded}"
        except Exception as e:
            logger.warning(f"UTF-8 인코딩 실패, ASCII 파일명 사용: {e}")
            # ASCII 파일명으로 fallback
            ascii_filename = f"TCFD_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            content_disposition = f"attachment; filename={ascii_filename}"
        
        # StreamingResponse로 반환 (메모리에서 직접)
        response = StreamingResponse(
            io.BytesIO(doc_bytes.getvalue()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": content_disposition,
                "Cache-Control": "no-cache, no-store, must-revalidate",
                "Pragma": "no-cache",
                "Expires": "0",
                "X-Content-Type-Options": "nosniff",
                "X-Frame-Options": "DENY"
            }
        )
        
        logger.info(f"Word 문서 응답 전송: {filename}")
        return response
        
    except Exception as e:
        logger.error(f"Word 문서 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=f"Word 문서 생성 실패: {str(e)}")

@tcfdreport_router.post("/download/pdf")
async def download_tcfd_report_as_pdf(data: Dict[str, Any], background_tasks: BackgroundTasks):
    """TCFD 보고서를 PDF로 다운로드"""
    try:
        logger.info(f"PDF 다운로드 요청: {data.get('company_name', 'Unknown')}")
        
        # 메모리에서 직접 PDF 생성
        try:
            pdf_bytes = await _generate_pdf_in_memory(data)
            
            # 파일명 생성
            company_name = data.get('company_name', 'TCFD')
            if company_name == 'TCFD' and data.get('draft'):
                draft_content = data['draft']
                if '**회사명**:' in draft_content:
                    company_name = draft_content.split('**회사명**:')[1].split('\n')[0].strip()
                elif '회사명:' in draft_content:
                    company_name = draft_content.split('회사명:')[1].split('\n')[0].strip()
            
            safe_company_name = company_name.replace('*', '').replace('/', '_').replace('\\', '_').replace(':', '_').replace('|', '_').replace('<', '_').replace('>', '_').replace('"', '_').replace('?', '_')
            if len(safe_company_name) > 20:
                safe_company_name = safe_company_name[:20]
            
            filename = f"{safe_company_name}_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            filename_encoded = urllib.parse.quote(filename)
            
            logger.info(f"PDF 생성 성공: {filename}, size={len(pdf_bytes)}B")
            
            # StreamingResponse로 반환 (메모리에서 직접)
            response = StreamingResponse(
                io.BytesIO(pdf_bytes),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename*=UTF-8''{filename_encoded}",
                    "Cache-Control": "no-cache, no-store, must-revalidate",
                    "Pragma": "no-cache",
                    "Expires": "0",
                    "X-Content-Type-Options": "nosniff",
                    "X-Frame-Options": "DENY"
                }
            )
            
            logger.info(f"PDF 응답 전송: {filename}")
            return response
            
        except Exception as e:
            logger.warning(f"PDF 생성 실패, HTML fallback 반환: {e}")
            return await _return_html_fallback(data, "weasyprint_error")
            
    except Exception as e:
        logger.error(f"PDF 다운로드 실패: {e}")
        raise HTTPException(status_code=500, detail=f"PDF 다운로드 실패: {str(e)}")

@tcfdreport_router.post("/download/combined")
async def download_tcfd_report_combined(data: Dict[str, Any], background_tasks: BackgroundTasks):
    """TCFD 보고서를 Word와 PDF로 생성하여 ZIP 파일로 다운로드"""
    try:
        logger.info(f"Combined 다운로드 요청: {data.get('company_name', 'Unknown')}")
        
        # 1. Word 문서 생성
        word_doc = await _create_word_document(data)
        if not word_doc:
            raise HTTPException(status_code=500, detail="Word 문서 생성 실패")
        
        # 2. PDF 생성 (메모리에서 직접)
        try:
            pdf_bytes = await _generate_pdf_in_memory(data)
            pdf_generated = True
        except Exception as e:
            logger.warning(f"PDF 생성 실패, HTML만 포함: {e}")
            pdf_bytes = None
            pdf_generated = False
        
        # 3. ZIP 파일 생성 (메모리에서)
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Word 문서 추가
            word_filename = f"{data.get('company_name', 'Unknown')}_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            zip_file.writestr(word_filename, word_doc.getvalue())
            
            # PDF 추가 (생성된 경우)
            if pdf_generated and pdf_bytes:
                pdf_filename = f"{data.get('company_name', 'Unknown')}_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                zip_file.writestr(pdf_filename, pdf_bytes)
            else:
                # HTML fallback 추가
                html_content = await _generate_html_content(data)
                html_filename = f"{data.get('company_name', 'Unknown')}_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
                zip_file.writestr(html_filename, html_content)
        
        zip_buffer.seek(0)
        
        # 4. 응답 반환
        filename = f"{data.get('company_name', 'Unknown')}_TCFD_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        filename_encoded = urllib.parse.quote(filename)
        
        response = StreamingResponse(
            io.BytesIO(zip_buffer.getvalue()),
            media_type="application/zip",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{filename_encoded}",
                "Cache-Control": "no-cache, no-store, must-revalidate",
                "Pragma": "no-cache",
                "Expires": "0",
                "X-Content-Type-Options": "nosniff",
                "X-Frame-Options": "DENY"
            }
        )
        
        logger.info(f"ZIP 파일 생성 완료: {filename}, size={len(zip_buffer.getvalue())}B")
        return response
        
    except Exception as e:
        logger.error(f"ZIP 파일 생성 실패: {e}")
        raise HTTPException(status_code=500, detail=f"ZIP 파일 생성 실패: {str(e)}")

async def _create_word_document(data: Dict[str, Any]) -> io.BytesIO:
    """Word 문서 생성 (메모리에서 직접)"""
    try:
        company_name = data.get('company_name', 'TCFD')
        if company_name == 'TCFD' and data.get('draft'):
            draft_content = data['draft']
            if '**회사명**:' in draft_content:
                company_name = draft_content.split('**회사명**:')[1].split('\n')[0].strip()
            elif '회사명:' in draft_content:
                company_name = draft_content.split('회사명:')[1].split('\n')[0].strip()
        
        safe_company_name = company_name.replace('*', '').replace('/', '_').replace('\\', '_').replace(':', '_').replace('|', '_').replace('<', '_').replace('>', '_').replace('"', '_').replace('?', '_')
        if len(safe_company_name) > 20:
            safe_company_name = safe_company_name[:20]
        
        doc = Document()
        
        title = doc.add_heading(f'{company_name} TCFD 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'생성일시: {datetime.now().strftime("%Y년 %m월 %d일 %H시 %M분")}')
        doc.add_paragraph('')
        
        doc.add_heading('AI 생성 초안', level=1)
        doc.add_paragraph(data.get('draft', ''))
        
        doc.add_heading('윤문된 텍스트', level=1)
        doc.add_paragraph(data.get('polished', ''))
        
        filename = f"{safe_company_name}_보고서_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info(f"Word 문서 생성 성공 (메모리): {filename}, size={len(doc_bytes.getvalue())}B")
        return doc_bytes
        
    except Exception as e:
        logger.error(f"Word 문서 생성 실패 (메모리): {e}")
        return None

async def _generate_pdf_in_memory(data: Dict[str, Any]) -> bytes:
    """메모리에서 직접 PDF 생성 (임시 파일 없이)"""
    try:
        # HTML 콘텐츠 생성
        html_content = await _generate_html_content(data)
        
        # CSS 스타일 정의
        css_content = """
        @page {
            size: A4;
            margin: 2cm;
            @top-center {
                content: "TCFD 기후 관련 재무정보 공시 보고서";
                font-size: 10pt;
                color: #666;
            }
            @bottom-center {
                content: counter(page);
                font-size: 10pt;
                color: #666;
            }
        }
        body {
            font-family: "Noto Sans CJK KR", "Nanum Gothic", Arial, sans-serif;
            font-size: 12pt;
            line-height: 1.6;
            color: #333;
        }
        h1, h2, h3 {
            color: #2c5aa0;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }
        h1 { font-size: 18pt; }
        h2 { font-size: 16pt; }
        h3 { font-size: 14pt; }
        .company-info {
            background-color: #f8f9fa;
            padding: 1em;
            border-left: 4px solid #2c5aa0;
            margin: 1em 0;
        }
        .section {
            margin: 1.5em 0;
            page-break-inside: avoid;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 1em 0;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 0.5em;
            text-align: left;
        }
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        """
        
        # WeasyPrint로 PDF 생성 (메모리에서)
        html = HTML(string=html_content, base_url=os.getcwd())
        css = CSS(string=css_content)
        
        pdf_bytes = html.write_pdf(stylesheets=[css])
        logger.info(f"PDF 생성 성공 (메모리): {len(pdf_bytes)}B")
        return pdf_bytes
        
    except Exception as e:
        logger.error(f"PDF 생성 실패: {e}")
        raise e

async def _generate_html_content(data: Dict[str, Any]) -> str:
    """HTML 콘텐츠 생성"""
    
    # draft와 polished 내용에서 TCFD 섹션 추출
    draft_content = data.get('draft', '')
    polished_content = data.get('polished', '')
    
    # TCFD 섹션별로 내용 파싱
    def extract_tcfd_section(content: str, section_keywords: list) -> str:
        """TCFD 섹션 내용 추출"""
        if not content:
            return '내용이 없습니다.'
        
        # 섹션 키워드로 내용 찾기
        for keyword in section_keywords:
            if keyword in content:
                # 해당 섹션부터 다음 섹션까지 추출
                start_idx = content.find(keyword)
                if start_idx != -1:
                    # 다음 섹션 찾기
                    next_section = None
                    for next_keyword in ['## 2.', '## 3.', '## 4.', '## 5.', '## 결론', '---']:
                        next_idx = content.find(next_keyword, start_idx + 1)
                        if next_idx != -1:
                            next_section = next_idx
                            break
                    
                    if next_section:
                        section_content = content[start_idx:next_section].strip()
                    else:
                        section_content = content[start_idx:].strip()
                    
                    # 마크다운 제거 및 정리
                    section_content = section_content.replace('##', '').replace('###', '').replace('**', '').replace('*', '').strip()
                    if section_content:
                        return section_content
        
        return '내용이 없습니다.'
    
    # 각 TCFD 섹션 내용 추출
    governance_content = extract_tcfd_section(draft_content, ['## 1. 거버넌스', '### G1:', '### G2:'])
    strategy_content = extract_tcfd_section(draft_content, ['## 2. 전략', '### S1:', '### S2:', '### S3:'])
    risk_management_content = extract_tcfd_section(draft_content, ['## 3. 위험 관리', '### R1:', '### R2:', '### R3:'])
    metrics_targets_content = extract_tcfd_section(draft_content, ['## 4. 지표 및 목표', '### M1:', '### M2:', '### M3:'])
    
    # 회사명 추출
    company_name = data.get('company_name', 'N/A')
    if company_name == 'N/A' and draft_content:
        if '**회사명**:' in draft_content:
            company_name = draft_content.split('**회사명**:')[1].split('\n')[0].strip()
        elif '회사명:' in draft_content:
            company_name = draft_content.split('회사명:')[1].split('\n')[0].strip()
    
    # 보고 연도 추출
    report_year = data.get('report_year', 'N/A')
    if report_year == 'N/A' and draft_content:
        if '**보고서 연도**:' in draft_content:
            report_year = draft_content.split('**보고서 연도**:')[1].split('\n')[0].strip()
        elif '보고서 연도:' in draft_content:
            report_year = draft_content.split('보고서 연도:')[1].split('\n')[0].strip()
    
    html_content = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>TCFD 기후 관련 재무정보 공시 보고서</title>
    </head>
    <body>
        <div class="company-info">
            <h1>TCFD 기후 관련 재무정보 공시 보고서</h1>
            <p><strong>기업명:</strong> {company_name}</p>
            <p><strong>보고 연도:</strong> {report_year}</p>
            <p><strong>생성 일시:</strong> {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}</p>
        </div>
        
        <div class="section">
            <h2>1. 거버넌스 (Governance)</h2>
            <p>{governance_content}</p>
        </div>
        
        <div class="section">
            <h2>2. 전략 (Strategy)</h2>
            <p>{strategy_content}</p>
        </div>
        
        <div class="section">
            <h2>3. 위험 관리 (Risk Management)</h2>
            <p>{risk_management_content}</p>
        </div>
        
        <div class="section">
            <h2>4. 지표 및 목표 (Metrics and Targets)</h2>
            <p>{metrics_targets_content}</p>
        </div>
        
        <div class="section">
            <h2>5. AI 생성 보고서 전문</h2>
            <h3>초안 (Draft)</h3>
            <div style="background-color: #f8f9fa; padding: 1em; border-radius: 4px; margin: 1em 0;">
                <pre style="white-space: pre-wrap; font-family: inherit; margin: 0;">{draft_content[:1000]}{'...' if len(draft_content) > 1000 else ''}</pre>
            </div>
            
            <h3>윤문된 텍스트 (Polished)</h3>
            <div style="background-color: #e8f5e8; padding: 1em; border-radius: 4px; margin: 1em 0;">
                <pre style="white-space: pre-wrap; font-family: inherit; margin: 0;">{polished_content[:1000]}{'...' if len(polished_content) > 1000 else ''}</pre>
            </div>
        </div>
    </body>
    </html>
    """
    return html_content