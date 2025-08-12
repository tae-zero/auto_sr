"""
TCFD Analysis Service - LangChain 기반 문서 분석
"""
import os
import logging
import tempfile
import shutil
from typing import Dict, List, Optional, Any
from fastapi import UploadFile

# LangChain imports
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate

logger = logging.getLogger(__name__)

class TCFDAnalysisService:
    """TCFD 문서 분석 서비스"""
    
    def __init__(self):
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.openai_model = os.getenv("OPENAI_MODEL", "gpt-4")
        self.vector_store = None
        self.conversation_chain = None
        self.embeddings = None
        
    async def initialize_ai_services(self) -> bool:
        """AI 서비스 초기화"""
        if not self.openai_api_key:
            logger.warning("⚠️ OPENAI_API_KEY가 설정되지 않았습니다.")
            return False
            
        try:
            # OpenAI 모델 초기화
            self.llm = ChatOpenAI(
                openai_api_key=self.openai_api_key,
                model_name=self.openai_model,
                temperature=0.3,
                max_tokens=2000
            )
            
            # 임베딩 모델 초기화
            self.embeddings = OpenAIEmbeddings(openai_api_key=self.openai_api_key)
            
            # 벡터 데이터베이스 초기화
            persist_directory = os.getenv("CHROMA_PERSIST_DIRECTORY", "./chroma_db")
            if os.path.exists(persist_directory):
                self.vector_store = Chroma(
                    persist_directory=persist_directory,
                    embedding_function=self.embeddings
                )
                logger.info("✅ 기존 벡터 데이터베이스 로드 완료")
            else:
                # 빈 벡터 스토어 생성
                self.vector_store = Chroma(
                    embedding_function=self.embeddings,
                    persist_directory=persist_directory
                )
                logger.info("✅ 새로운 벡터 데이터베이스 생성 완료")
            
            # TCFD 전용 프롬프트 템플릿
            tcfd_prompt_template = """당신은 TCFD(기후 관련 재무정보 공개) 전문가입니다.
            
            다음 컨텍스트를 사용하여 질문에 답변하세요:
            {context}
            
            질문: {question}
            
            TCFD 프레임워크의 다음 4개 영역을 고려하여 답변하세요:
            1. 거버넌스 (Governance)
            2. 전략 (Strategy)
            3. 리스크 관리 (Risk Management)
            4. 지표 및 목표 (Metrics and Targets)
            
            답변:"""
            
            prompt = PromptTemplate(
                template=tcfd_prompt_template,
                input_variables=["context", "question"]
            )
            
            # 대화 체인 초기화
            memory = ConversationBufferMemory(
                memory_key="chat_history",
                return_messages=True
            )
            
            self.conversation_chain = ConversationalRetrievalChain.from_llm(
                llm=self.llm,
                retriever=self.vector_store.as_retriever(),
                memory=memory,
                return_source_documents=True,
                combine_docs_chain_kwargs={"prompt": prompt}
            )
            
            logger.info("✅ TCFD AI 서비스 초기화 완료")
            return True
            
        except Exception as e:
            logger.error(f"❌ TCFD AI 서비스 초기화 실패: {str(e)}")
            return False
    
    async def analyze_report(self, file: UploadFile, company_info: Dict[str, Any]) -> Dict[str, Any]:
        """TCFD 보고서 AI 분석"""
        try:
            if not self.vector_store:
                raise Exception("AI 서비스가 초기화되지 않았습니다.")
            
            # 문서 업로드 및 벡터화
            doc_info = await self.upload_document(file, company_info.get("company_id", "unknown"), "tcfd_report")
            
            # TCFD 분석 프롬프트
            analysis_prompt = f"""
            회사 정보: {company_info}
            
            다음 TCFD 보고서를 분석하여 다음 사항을 평가하세요:
            
            1. 거버넌스 (Governance)
               - 기후 리스크 관리 책임자
               - 이사회 감독 체계
            
            2. 전략 (Strategy)
               - 기후 리스크 및 기회 식별
               - 시나리오 분석
               - 재무 영향
            
            3. 리스크 관리 (Risk Management)
               - 기후 리스크 통합 관리
               - 리스크 평가 프로세스
            
            4. 지표 및 목표 (Metrics and Targets)
               - GHG 배출량
               - 기후 관련 목표
            
            분석 결과를 구조화된 JSON 형태로 제공하세요.
            """
            
            # AI 분석 수행
            analysis_result = await self._perform_ai_analysis(analysis_prompt)
            
            return {
                "analysis_id": doc_info["document_id"],
                "company_info": company_info,
                "analysis_result": analysis_result,
                "uploaded_document": doc_info,
                "analysis_timestamp": "2024-01-01T00:00:00Z"  # 실제로는 현재 시간
            }
            
        except Exception as e:
            logger.error(f"TCFD 보고서 분석 실패: {str(e)}")
            raise
    
    async def upload_document(self, file: UploadFile, company_id: str, document_type: str) -> Dict[str, Any]:
        """문서 업로드 및 벡터화"""
        try:
            if not self.vector_store:
                raise Exception("AI 서비스가 초기화되지 않았습니다.")
            
            # 파일 크기 검증
            max_size = int(os.getenv("MAX_FILE_SIZE", "20971520"))  # 20MB
            if file.size > max_size:
                raise Exception(f"파일 크기가 너무 큽니다. 최대 {max_size/1024/1024}MB")
            
            # 파일 타입 검증
            supported_types = os.getenv("SUPPORTED_FILE_TYPES", "txt,pdf,docx,md").split(",")
            file_extension = file.filename.split(".")[-1].lower()
            if file_extension not in supported_types:
                raise Exception(f"지원하지 않는 파일 타입입니다. 지원: {', '.join(supported_types)}")
            
            # 임시 파일로 저장
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
                shutil.copyfileobj(file.file, temp_file)
                temp_path = temp_file.name
            
            try:
                # 문서 내용 추출 및 벡터화
                documents = await self._process_document(temp_path, file_extension, company_id, document_type)
                
                # 벡터 데이터베이스에 저장
                self.vector_store.add_documents(documents)
                self.vector_store.persist()
                
                return {
                    "document_id": f"doc_{company_id}_{document_type}_{len(documents)}",
                    "filename": file.filename,
                    "company_id": company_id,
                    "document_type": document_type,
                    "documents_count": len(documents),
                    "file_size": file.size,
                    "file_type": file_extension
                }
                
            finally:
                # 임시 파일 삭제
                os.unlink(temp_path)
                
        except Exception as e:
            logger.error(f"문서 업로드 실패: {str(e)}")
            raise
    
    async def search_knowledge(self, query: str, company_id: Optional[str] = None, top_k: int = 5) -> List[Dict[str, Any]]:
        """TCFD 지식 베이스 검색"""
        try:
            if not self.vector_store:
                raise Exception("AI 서비스가 초기화되지 않았습니다.")
            
            # 검색 쿼리에 회사 정보 추가
            if company_id:
                enhanced_query = f"회사 ID: {company_id}, {query}"
            else:
                enhanced_query = query
            
            # 벡터 검색 수행
            results = self.vector_store.similarity_search(enhanced_query, k=top_k)
            
            search_results = []
            for doc in results:
                search_results.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "relevance_score": 0.95  # 실제로는 유사도 점수 계산
                })
            
            return search_results
            
        except Exception as e:
            logger.error(f"지식 검색 실패: {str(e)}")
            raise
    
    async def _process_document(self, file_path: str, file_type: str, company_id: str, document_type: str) -> List[Document]:
        """문서 처리 및 텍스트 추출"""
        try:
            if file_type == "txt":
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_type == "md":
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_type == "pdf":
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
            elif file_type == "docx":
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
            else:
                raise ValueError(f"지원하지 않는 파일 타입: {file_type}")
            
            # 텍스트 분할
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            )
            
            chunks = text_splitter.split_text(text)
            
            # Document 객체 생성
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "company_id": company_id,
                        "document_type": document_type,
                        "chunk_id": i,
                        "file_type": file_type
                    }
                )
                documents.append(doc)
            
            return documents
            
        except Exception as e:
            logger.error(f"문서 처리 실패: {str(e)}")
            raise
    
    async def _perform_ai_analysis(self, prompt: str) -> Dict[str, Any]:
        """AI 분석 수행"""
        try:
            if not self.conversation_chain:
                raise Exception("AI 서비스가 초기화되지 않았습니다.")
            
            # AI 분석 실행
            response = self.conversation_chain.invoke({
                "question": prompt,
                "chat_history": []
            })
            
            return {
                "analysis": response["answer"],
                "sources": [doc.page_content[:100] + "..." for doc in response.get("source_documents", [])],
                "model_used": self.openai_model
            }
            
        except Exception as e:
            logger.error(f"AI 분석 수행 실패: {str(e)}")
            raise
